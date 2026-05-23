"""Extract structured text from form_solution_solve2026.docx."""
from pathlib import Path
from docx import Document

src = Path(r"D:\projects\LegalShield\docs\grants\ristex_solve_2026\form_solution_solve2026.docx")
doc = Document(str(src))

out_lines = []
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ""
    t = p.text.strip()
    if t:
        out_lines.append(f"[{style}] {t}")

# Tables
for ti, t in enumerate(doc.tables):
    out_lines.append(f"\n=== TABLE {ti} ({len(t.rows)}x{len(t.columns)}) ===")
    for ri, row in enumerate(t.rows):
        cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
        out_lines.append(f"  R{ri}: " + " || ".join(cells))

out = src.with_suffix(".txt")
out.write_text("\n".join(out_lines), encoding="utf-8")
print(f"paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)} -> {out.name} ({out.stat().st_size} bytes)")
