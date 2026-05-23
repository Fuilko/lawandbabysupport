"""
Litigation Document Text Extractor
===================================

CALL4 と 裁判所書式集 で集めた PDF / DOCX / XLSX から全文テキストを抽出し、
DuckDB の `litigation_chunks` テーブルに保存する。

抽出ポリシー:
  - PDF: pdfplumber でページ単位抽出
  - DOCX: python-docx で段落 + 表抽出
  - XLSX: openpyxl で各シート抽出
  - 失敗ファイルは sha256 とともに errors テーブルに記録

チャンキング戦略（RAG 向け）:
  - 512 字を 1 chunk、50 字オーバーラップ
  - 文書境界、段落境界を尊重
  - 各 chunk に source_id, chunk_idx, char_start/end を保持

使い方:
  python -m legalshield.crawlers.litigation_docs.extract_text
"""
from __future__ import annotations

import hashlib
import logging
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

import duckdb

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
REPO_ROOT = Path(__file__).resolve().parents[3]
DB_PATH = REPO_ROOT / "legalshield" / "lancedb" / "litigation.duckdb"

CHUNK_SIZE = 512
CHUNK_OVERLAP = 50

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger("extract_text")


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------
def extract_pdf(path: Path) -> str:
    import pdfplumber
    out: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages):
            try:
                t = page.extract_text() or ""
            except Exception as e:
                log.debug("pdf page %d extract fail: %s", i, e)
                t = ""
            if t:
                out.append(t)
    return "\n\n".join(out)


def extract_docx(path: Path) -> str:
    import docx
    doc = docx.Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(path), data_only=True, read_only=True)
    parts: list[str] = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        parts.append(f"# Sheet: {sheet}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


EXTRACTORS = {
    "pdf": extract_pdf,
    "docx": extract_docx,
    "xlsx": extract_xlsx,
}


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------
def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[tuple[int, int, str]]:
    """Return list of (char_start, char_end, chunk_text). Paragraph-aware."""
    if not text:
        return []
    chunks: list[tuple[int, int, str]] = []
    n = len(text)
    start = 0
    while start < n:
        end = min(start + size, n)
        # Try to break at a paragraph or sentence boundary
        if end < n:
            for sep in ["\n\n", "\n", "。", "．", "."]:
                idx = text.rfind(sep, start + size // 2, end)
                if idx != -1:
                    end = idx + len(sep)
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append((start, end, chunk))
        if end >= n:
            break
        start = end - overlap
        if start < 0:
            start = 0
    return chunks


# ---------------------------------------------------------------------------
# Schema
# ---------------------------------------------------------------------------
DDL = """
CREATE TABLE IF NOT EXISTS litigation_chunks (
    chunk_id        VARCHAR PRIMARY KEY,   -- {source_type}__{source_id}__{idx}
    source_type     VARCHAR,               -- 'call4' or 'court_form'
    source_id       VARCHAR,               -- case_id or form_id
    chunk_idx       INTEGER,
    char_start      INTEGER,
    char_end        INTEGER,
    text            TEXT,
    text_sha256     VARCHAR,
    category        VARCHAR,               -- 法律分類
    title           VARCHAR,
    source_url      VARCHAR,
    extracted_at    VARCHAR
);

CREATE TABLE IF NOT EXISTS extraction_errors (
    source_type     VARCHAR,
    source_id       VARCHAR,
    file_path       VARCHAR,
    error           TEXT,
    failed_at       VARCHAR
);
"""


# ---------------------------------------------------------------------------
# Main processors
# ---------------------------------------------------------------------------
def process_call4(con: duckdb.DuckDBPyConnection) -> int:
    """Chunk CALL4 case full_text already in DB."""
    log.info("Processing CALL4 cases...")
    rows = con.execute(
        "SELECT case_id, title_ja, url, full_text FROM call4_cases WHERE full_text IS NOT NULL"
    ).fetchall()
    log.info("  %d CALL4 cases to chunk", len(rows))

    # Wipe and rebuild only call4 entries to keep idempotent
    con.execute("DELETE FROM litigation_chunks WHERE source_type = 'call4'")
    count = 0
    for case_id, title, url, full_text in rows:
        chunks = chunk_text(full_text or "")
        for idx, (s, e, txt) in enumerate(chunks):
            chunk_id = f"call4__{case_id}__{idx:04d}"
            sha = hashlib.sha256(txt.encode("utf-8")).hexdigest()
            con.execute(
                "INSERT INTO litigation_chunks VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
                [
                    chunk_id, "call4", case_id, idx, s, e, txt, sha,
                    "公共訴訟", title, url, _now_iso(),
                ],
            )
            count += 1
    log.info("  → %d CALL4 chunks written", count)
    return count


def process_court_forms(con: duckdb.DuckDBPyConnection) -> int:
    """Extract text from court form files and chunk."""
    log.info("Processing court forms...")
    rows = con.execute(
        "SELECT form_id, category, title, file_url, file_ext, local_path FROM court_forms"
    ).fetchall()
    log.info("  %d court forms to extract", len(rows))

    con.execute("DELETE FROM litigation_chunks WHERE source_type = 'court_form'")

    count = 0
    err_count = 0
    for form_id, category, title, file_url, file_ext, local_path in rows:
        ext = (file_ext or "").lower()
        extractor = EXTRACTORS.get(ext)
        if extractor is None:
            continue
        full_path = REPO_ROOT / local_path
        if not full_path.exists():
            continue
        try:
            text = extractor(full_path)
        except Exception as e:
            log.warning("  EXTRACT FAIL %s (%s): %s", form_id, ext, e)
            con.execute(
                "INSERT INTO extraction_errors VALUES (?,?,?,?,?)",
                ["court_form", form_id, str(local_path), str(e), _now_iso()],
            )
            err_count += 1
            continue
        chunks = chunk_text(text)
        for idx, (s, e, txt) in enumerate(chunks):
            chunk_id = f"court_form__{form_id}__{idx:04d}"
            sha = hashlib.sha256(txt.encode("utf-8")).hexdigest()
            con.execute(
                "INSERT INTO litigation_chunks VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
                [
                    chunk_id, "court_form", form_id, idx, s, e, txt, sha,
                    category, title, file_url, _now_iso(),
                ],
            )
            count += 1
    log.info("  → %d court_form chunks written (errors: %d)", count, err_count)
    return count


def run() -> dict:
    log.info("============================================================")
    log.info("Litigation Text Extractor — start (DB: %s)", DB_PATH)
    log.info("============================================================")
    con = duckdb.connect(str(DB_PATH))
    con.execute(DDL)

    n_call4 = process_call4(con)
    n_forms = process_court_forms(con)

    total = con.execute("SELECT COUNT(*) FROM litigation_chunks").fetchone()[0]
    by_type = con.execute(
        "SELECT source_type, COUNT(*), SUM(LENGTH(text)) FROM litigation_chunks GROUP BY source_type"
    ).fetchall()
    con.close()

    summary = {
        "call4_chunks": n_call4,
        "court_form_chunks": n_forms,
        "total_chunks": total,
        "by_type": by_type,
    }
    log.info("Done: %s", summary)
    return summary


if __name__ == "__main__":
    run()
